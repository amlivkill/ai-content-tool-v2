import streamlit as st
import requests
import re
import PyPDF2
import docx
import csv
import io

# Mobile-friendly page config
st.set_page_config(
    page_title="Content Analyzer Pro",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS
st.markdown("""
<style>
    .main-header { font-size: 2rem; color: #1f77b4; text-align: center; }
    .feature-card { background: #f8f9fa; padding: 1rem; border-radius: 10px; margin: 0.5rem 0; }
    .file-badge { background: #1f77b4; color: white; padding: 0.2rem 0.6rem; border-radius: 10px; font-size: 0.8rem; }
    @media (max-width: 768px) {
        .main-header { font-size: 1.5rem !important; }
        .stButton button, .stDownloadButton button { width: 100% !important; }
    }
</style>
""", unsafe_allow_html=True)

def simple_scrape_website(url):
    """BeautifulSoup के बिना website scraping"""
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, timeout=10)
        
        # Simple title extraction
        title_match = re.search(r'<title[^>]*>(.*?)</title>', response.text, re.IGNORECASE)
        title = title_match.group(1).strip() if title_match else "No Title"
        
        # Simple content cleaning
        clean_text = re.sub(r'<[^>]+>', ' ', response.text)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        
        return {
            'success': True,
            'title': title[:50],
            'content': clean_text[:3000],
            'url': url,
            'content_length': len(clean_text)
        }
    except Exception as e:
        return {'success': False, 'error': str(e), 'url': url}

def extract_pdf_content(file):
    """PDF से text extract"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return {
            'success': True,
            'content': text,
            'pages': len(pdf_reader.pages),
            'characters': len(text),
            'words': len(text.split())
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def extract_docx_content(file):
    """DOCX से text extract"""
    try:
        doc = docx.Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return {
            'success': True,
            'content': text,
            'paragraphs': len(doc.paragraphs),
            'characters': len(text),
            'words': len(text.split())
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def extract_csv_content(file):
    """CSV analysis without pandas"""
    try:
        content = file.read().decode("utf-8")
        reader = csv.reader(io.StringIO(content))
        rows = list(reader)
        
        if not rows:
            return {'success': False, 'error': "Empty CSV file"}
        
        return {
            'success': True,
            'rows': rows,
            'row_count': len(rows),
            'column_count': len(rows[0]) if rows else 0,
            'headers': rows[0] if rows else []
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def extract_txt_content(file):
    """TXT file analysis"""
    try:
        content = file.read().decode("utf-8")
        return {
            'success': True,
            'content': content,
            'characters': len(content),
            'words': len(content.split()),
            'lines': len(content.splitlines())
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def main():
    st.markdown('<h1 class="main-header">🔍 Content Analyzer Pro</h1>', unsafe_allow_html=True)
    st.markdown("**🌐 URLs + 📄 PDF + 📝 DOCX + 📊 CSV + 📃 TXT - All in One**")
    
    # Features
    col1, col2, col3, col4, col5 = st.columns(5)
    formats = ["🌐 URLs", "📄 PDF", "📝 DOCX", "📊 CSV", "📃 TXT"]
    for i, format in enumerate(formats):
        with [col1, col2, col3, col4, col5][i]:
            st.markdown(f'<div class="feature-card" style="text-align:center;"><h4>{format}</h4></div>', unsafe_allow_html=True)
    
    # Tabs
    tab1, tab2 = st.tabs(["🌐 URL Scraping", "📁 File Analysis"])
    
    with tab1:
        st.header("URLs से Content Scrape करें")
        url_input = st.text_area("URLs डालें (एक line में एक)", placeholder="https://example.com", height=80)
        
        if st.button("🚀 Scrape URLs", use_container_width=True) and url_input:
            urls = [url.strip() for url in url_input.split('\n') if url.strip()]
            for i, url in enumerate(urls):
                with st.spinner(f"Scraping {url[:30]}..."):
                    result = simple_scrape_website(url)
                    if result['success']:
                        with st.expander(f"✅ {result['title']}"):
                            st.write(f"**URL:** {result['url']}")
                            st.write(f"**Length:** {result['content_length']} chars")
                            st.text_area("Preview", result['content'][:400] + "...", height=120, key=f"url_{i}")
                            st.download_button(
                                "📥 Download", result['content'], f"scraped_{i+1}.txt", 
                                "text/plain", use_container_width=True, key=f"dl_{i}"
                            )
                    else:
                        st.error(f"❌ {url} - {result['error']}")
    
    with tab2:
        st.header("Files Upload और Analyze करें")
        uploaded_files = st.file_uploader(
            "PDF, DOCX, CSV, TXT files select करें",
            type=['pdf', 'docx', 'csv', 'txt'],
            accept_multiple_files=True
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} files uploaded!")
            
            for file in uploaded_files:
                badge = f'<span class="file-badge">{file.type}</span>'
                st.markdown(badge, unsafe_allow_html=True)
                
                with st.expander(f"📄 {file.name}", expanded=True):
                    st.write(f"**Size:** {file.size / 1024:.1f} KB")
                    
                    if st.button(f"Analyze {file.name}", key=file.name, use_container_width=True):
                        # PDF
                        if file.type == "application/pdf":
                            result = extract_pdf_content(file)
                            if result['success']:
                                st.success("✅ PDF Analyzed!")
                                st.write(f"**Pages:** {result['pages']}")
                                st.write(f"**Words:** {result['words']}")
                                st.text_area("Content", result['content'][:500] + "...", height=150, key=f"pdf_{file.name}")
                                st.download_button("📥 Download", result['content'], f"{file.name}.txt", "text/plain", use_container_width=True)
                            else:
                                st.error(f"❌ {result['error']}")
                        
                        # DOCX
                        elif "docx" in file.type:
                            result = extract_docx_content(file)
                            if result['success']:
                                st.success("✅ DOCX Analyzed!")
                                st.write(f"**Paragraphs:** {result['paragraphs']}")
                                st.write(f"**Words:** {result['words']}")
                                st.text_area("Content", result['content'][:500] + "...", height=150, key=f"docx_{file.name}")
                                st.download_button("📥 Download", result['content'], f"{file.name}.txt", "text/plain", use_container_width=True)
                            else:
                                st.error(f"❌ {result['error']}")
                        
                        # CSV
                        elif file.type == "text/csv":
                            result = extract_csv_content(file)
                            if result['success']:
                                st.success("✅ CSV Analyzed!")
                                st.write(f"**Rows:** {result['row_count']}")
                                st.write(f"**Columns:** {result['column_count']}")
                                st.write(f"**Headers:** {', '.join(result['headers'])}")
                                
                                # Show first 5 rows
                                st.write("**First 5 Rows:**")
                                for i, row in enumerate(result['rows'][:5]):
                                    st.write(f"{i+1}. {row}")
                                
                                # Download as CSV
                                csv_content = "\n".join([",".join(row) for row in result['rows']])
                                st.download_button("📥 Download CSV", csv_content, f"analyzed_{file.name}", "text/csv", use_container_width=True)
                            else:
                                st.error(f"❌ {result['error']}")
                        
                        # TXT
                        elif file.type == "text/plain":
                            result = extract_txt_content(file)
                            if result['success']:
                                st.success("✅ Text Analyzed!")
                                st.write(f"**Characters:** {result['characters']}")
                                st.write(f"**Words:** {result['words']}")
                                st.write(f"**Lines:** {result['lines']}")
                                st.text_area("Content", result['content'][:500] + "...", height=150, key=f"txt_{file.name}")
                                st.download_button("📥 Download", result['content'], f"analyzed_{file.name}", "text/plain", use_container_width=True)
                            else:
                                st.error(f"❌ {result['error']}")

if __name__ == "__main__":
    main()
